from io import BytesIO
from unittest.mock import patch, MagicMock, PropertyMock

import pytest
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from src.alita_tools.sharepoint.utils import read_docx_from_bytes


# --- Fixtures ---

@pytest.fixture
def valid_docx_bytes():
    """Fixture to create valid .docx bytes."""
    doc = Document()
    doc.add_paragraph("Paragraph 1")
    doc.add_paragraph("Paragraph 2")
    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream.getvalue()

@pytest.fixture
def empty_docx_bytes():
    """Fixture to create empty .docx bytes."""
    doc = Document()
    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream.getvalue()

# --- Test Class ---

@pytest.mark.unit
@pytest.mark.sharepoint
@pytest.mark.utils
class TestSharepointReadDocxFromBytes:

    @pytest.mark.positive
    def test_read_docx_from_bytes_success(self, valid_docx_bytes):
        """Test successful reading of .docx content from bytes."""
        result = read_docx_from_bytes(valid_docx_bytes)
        assert result == "Paragraph 1\nParagraph 2"

    @pytest.mark.positive
    def test_read_docx_from_bytes_empty_file(self, empty_docx_bytes):
        """Test behavior with a structurally valid but empty .docx file."""
        result = read_docx_from_bytes(empty_docx_bytes)
        assert result == "" # No paragraphs, so empty string

    @pytest.mark.negative
    def test_read_docx_from_bytes_invalid_format(self, caplog):
        """Test behavior with bytes that are not a valid docx file."""
        invalid_content = b"This is not a valid docx file."
        result = read_docx_from_bytes(invalid_content)
        assert result == ""
        # Check the full log message format
        assert "Error reading .docx from bytes: file seems not to be a valid package" in caplog.text

    @pytest.mark.negative
    def test_read_docx_from_bytes_none_input(self, caplog):
        """Test behavior when input is None."""
        result = read_docx_from_bytes(None)
        assert result == ""
        # Check the full log message format
        assert "Error reading .docx from bytes" in caplog.text

    @pytest.mark.negative
    @patch('src.alita_tools.sharepoint.utils.Document')
    def test_read_docx_from_bytes_document_processing_error(self, mock_document_class, valid_docx_bytes, caplog):
        """Test behavior when accessing paragraphs fails."""
        mock_doc_instance = MagicMock()
        # Simulate an error when accessing paragraphs
        mock_doc_instance.paragraphs = PropertyMock(side_effect=Exception("Cannot read paragraphs"))
        mock_document_class.return_value = mock_doc_instance

        result = read_docx_from_bytes(valid_docx_bytes)

        assert result == ""
        # Check the full log message format
        assert "Error reading .docx from bytes: Cannot read paragraphs" in caplog.text
        mock_document_class.assert_called_once() # Ensure Document was called
import logging
import pytest
from unittest.mock import patch, MagicMock
from io import BytesIO
from docx import Document

from src.alita_tools.sharepoint.utils import read_docx_from_bytes

@pytest.mark.unit
@pytest.mark.sharepoint
class TestSharepointReadDocxFromBytes:
    
    @pytest.mark.positive
    def test_read_docx_from_bytes_success(self):
        """Test successful reading of a valid .docx file."""
        # Create a mock Document with paragraphs
        mock_doc = MagicMock(spec=Document)
        mock_paragraph1 = MagicMock()
        mock_paragraph1.text = "First paragraph"
        mock_paragraph2 = MagicMock()
        mock_paragraph2.text = "Second paragraph"
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]
        
        # Patch Document to return our mock
        with patch('src.alita_tools.sharepoint.utils.Document', return_value=mock_doc):
            result = read_docx_from_bytes(b'mock_docx_content')
        
        assert result == "First paragraph\nSecond paragraph"
    
    @pytest.mark.negative
    def test_read_docx_from_bytes_invalid_format(self, caplog):
        """Test reading an invalid .docx file format."""
        # Patch Document to raise an exception
        with patch('src.alita_tools.sharepoint.utils.Document', 
                  side_effect=Exception("file seems not to be a valid package")):
            with caplog.at_level(logging.ERROR):
                result = read_docx_from_bytes(b'invalid_docx_content')
        
        assert result == ""
        assert "Error reading .docx from bytes: file seems not to be a valid package" in caplog.text
    
    @pytest.mark.negative
    def test_read_docx_from_bytes_document_processing_error(self, caplog):
        """Test error during document processing after successful loading."""
        # Create a mock Document that raises when paragraphs is accessed
        mock_doc = MagicMock(spec=Document)
        
        # Set up the mock to raise an exception when paragraphs is accessed
        type(mock_doc).paragraphs = PropertyMock(side_effect=Exception("Cannot read paragraphs"))
        
        # Patch Document to return our problematic mock
        with patch('src.alita_tools.sharepoint.utils.Document', return_value=mock_doc):
            with caplog.at_level(logging.ERROR):
                result = read_docx_from_bytes(b'mock_docx_content')
        
        assert result == ""
        assert "Error reading .docx from bytes: Cannot read paragraphs" in caplog.text
